from pathlib import Path

import pdfplumber
from docx import Document


ROOT = Path("/Users/cj/careloop")
OUT = ROOT / "tmp/careloop_skywork/text"
OUT.mkdir(parents=True, exist_ok=True)


def extract_docx(source: Path, target: Path) -> None:
    doc = Document(source)
    lines: list[str] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"[P{index:03d} | {paragraph.style.name}] {text}")
    for table_index, table in enumerate(doc.tables, start=1):
        lines.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(" | ".join(cells))
    target.write_text("\n".join(lines), encoding="utf-8")


def extract_pdf(source: Path, target: Path) -> None:
    chunks: list[str] = []
    with pdfplumber.open(source) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            chunks.append(f"\n===== SLIDE {page_number:02d} =====\n{text.strip()}")
    target.write_text("\n".join(chunks), encoding="utf-8")


extract_docx(
    ROOT / "michk/Careloop_Maker_in_China_HK_2026_PPT_Optimization_Guide_v3.docx",
    OUT / "optimization_guide.txt",
)
extract_pdf(
    ROOT / "slides/TeaVita AI--Maker in China.pdf",
    OUT / "teavita_slides.txt",
)
extract_pdf(
    ROOT / "slides/Group 5-暖桌Careloop.pdf",
    OUT / "former_careloop_slides.txt",
)
