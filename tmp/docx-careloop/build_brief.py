from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

OUT = Path('/Users/cj/careloop/michk/Careloop_Maker_in_China_HK_2026_Demo_Improvement_Brief.docx')

BLUE = '1F4E79'
BLUE_DARK = '17365D'
TEAL = '0F6864'
CORAL = 'D85B49'
GRAY = '5B6573'
LIGHT = 'F2F4F7'
LIGHT_BLUE = 'E8EEF5'
PALE_TEAL = 'E6F2EF'
GOLD = 'B07A00'
RED = '9B1C1C'
BLACK = '1F1F1F'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in('w:tcW')
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def set_font(run, name='Calibri', size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rfonts.set(qn('w:eastAsia'), 'Songti SC')
    rfonts.set(qn('w:cs'), name)
    rfonts.set(qn('w:hint'), 'eastAsia')
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rpr.append(lang)
    lang.set(qn('w:eastAsia'), 'zh-CN')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instruction
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, size=9, color=GRAY)

def add_rule(paragraph, color=BLUE):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_para(doc, text='', style=None, size=11, color=BLACK, bold=False, italic=False, align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    if text:
        r = p.add_run(text)
        set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p

def add_rich_para(doc, parts, style=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    for text, kwargs in parts:
        r = p.add_run(text)
        set_font(r, **kwargs)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1:16, 2:13, 3:12}[level], color={1:BLUE, 2:BLUE, 3:BLUE_DARK}[level], bold=True)
    return p

def bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, size=11, color=BLACK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=BLACK)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r, size=11, color=BLACK)
    return p

def callout(doc, label, text, fill=PALE_TEAL, label_color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [9360])
    c = table.cell(0,0)
    set_cell_shading(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + '  ')
    set_font(r, size=11, color=label_color, bold=True)
    r = p.add_run(text)
    set_font(r, size=11, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(t, widths)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=font_size, color=BLUE_DARK, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(val))
            set_font(r, size=font_size, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('来源 / 备注：' + text)
    set_font(r, size=9, color=GRAY, italic=True)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = hp.add_run('CARELOOP  |  Maker in China Hong Kong 2026')
        set_font(r, size=9, color=GRAY, bold=True)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run('Internal working brief  |  ')
        set_font(r, size=9, color=GRAY)
        add_field(fp, 'PAGE')

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')
    normal._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
    lang = normal._element.rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        normal._element.rPr.append(lang)
    lang.set(qn('w:eastAsia'), 'zh-CN')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,BLUE_DARK,8,4)]:
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')
        st._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
        lang = st._element.rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            st._element.rPr.append(lang)
        lang.set(qn('w:eastAsia'), 'zh-CN')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ['List Bullet','List Bullet 2','List Number']:
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')
        st._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
        lang = st._element.rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            st._element.rPr.append(lang)
        lang.set(qn('w:eastAsia'), 'zh-CN')
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.10
    doc.styles['List Bullet'].paragraph_format.left_indent = Inches(0.5)
    doc.styles['List Bullet'].paragraph_format.first_line_indent = Inches(-0.25)
    doc.styles['List Bullet 2'].paragraph_format.left_indent = Inches(0.75)
    doc.styles['List Bullet 2'].paragraph_format.first_line_indent = Inches(-0.25)
    doc.styles['List Number'].paragraph_format.left_indent = Inches(0.5)
    doc.styles['List Number'].paragraph_format.first_line_indent = Inches(-0.25)

def build():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    # Cover / proposal centerpiece adapted to internal brief.
    add_para(doc, 'CARELOOP', size=12, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, '“创客中国”香港站 2026', size=25, color=BLUE_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, 'Demo 与参赛材料完整改进说明书', size=17, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, '基于旧版《Group 5-暖桌Careloop》、TeaVita AI 已提交样本、makerinchina.hk 公开信息及当前 Careloop project', size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    callout(doc, '核心判断', 'Careloop 不需要假装已经拥有市场收入或真实试点；在当前时间节点，新版应把“可运行的软件原型”转化为可信的技术证据，并把线下试点放到赛后路线图。', fill=PALE_TEAL)
    table(doc, ['文档信息', '内容'], [
        ['用途', '内部改稿、Demo 排练、报名材料准备'],
        ['目标受众', '香港站评审、投资者、社区/食堂/物业合作方'],
        ['事实基准', '截至 2026-08-20；所有未验证数据均须标注为假设'],
        ['当前产品状态', 'Web Demo v0.1.0；支付、AICAN、传送带为 Demo adapter'],
    ], [2100,7260], header_fill=LIGHT_BLUE, font_size=10)
    source(doc, '内部工作版本，不是对外提交稿。提交前请补齐团队主体、香港公司资格、访谈与合作方信息。')

    page_break(doc)
    heading(doc, '执行摘要', 1)
    callout(doc, '新版定位', '把 Careloop 从“社区自动化餐饮生态概念”改为“面向长者、家属、工作人员和社区食堂的适老智能点餐与履约操作系统”。', fill=LIGHT_BLUE, label_color=BLUE_DARK)
    heading(doc, '这次改稿的四个目标', 2)
    for x in [
        '符合香港站公开定位：香港非上市公司、智慧生活/智慧出行、人工智能与大数据、智能装备与机器人等科技领域。',
        '真实呈现当前产品：用新版 UI 做核心 Demo，清楚区分已实现、模拟适配器和未来研发。',
        '补足“空”的内容：用公开市场证据、技术验证证据、可远程完成的验证材料和赛后香港试点设计取代虚构 traction。',
        '提高评委记忆点：在短时间内让评委理解问题、产品、技术闭环、商业路径和下一步合作诉求。',
    ]:
        bullet(doc, x)
    heading(doc, '一句话版本', 2)
    add_para(doc, 'Careloop 是一个适老社区餐饮操作系统：让长者用语音或一键完成点餐，让家属可以远程代点，让工作人员可以接管异常订单，让社区食堂可以看见完整履约流程。', size=13, color=BLUE_DARK, bold=True, after=8)
    heading(doc, '事实边界', 2)
    table(doc, ['证据类型', '当前可说', '当前不可说'], [
        ['已实现', 'Web 点餐、语音预览、代点、人工协助、订单/后厨/配送状态流转', '真实支付、真实机器人、真实传送带'],
        ['工程验证', '2 个测试文件、9 个测试通过；build 与 lint 已通过', '真实用户成功率、真实复购率、真实订单量'],
        ['商业假设', '赛后香港社区试点；食堂/物业/家属三方价值', '10 个社区、8 个月回本、RMB 500,000 已足以规模化'],
        ['未来研发', '硬件自动化、营养推荐、社区运营数据', '已形成专利或数据护城河'],
    ], [1800,3780,3780], font_size=9.5)

    page_break(doc)
    heading(doc, '一、赛事适配：从学校创客展示转向香港站项目路演', 1)
    add_para(doc, '官网目前公开显示，2026 香港站欢迎香港非上市公司参加，覆盖智慧生活与智慧出行、人工智能与大数据、智能装备与机器人、生物医药与健康等领域；赛事目标包括推动创新技术与企业、产业园区、市场和政府政策对接。', after=8)
    table(doc, ['赛事公开信号', '对 Careloop 的含义', 'Slides 必须出现的证据'], [
        ['香港非上市公司参赛', '先确认主体资格与报名材料一致', '公司/团队主体、注册地、负责人和联系人'],
        ['技术与产业对接', '不能只讲情怀，要讲可部署产品', '系统闭环、实施条件、合作方画像'],
        ['香港→内地/大湾区连接', '香港应成为第一验证场景，但不必在提交前完成线下落地', '赛后香港试点，再讲大湾区复制'],
        ['十强有投资者、企业家和园区对接', '结尾要提出可执行的合作请求', '需要远程评审/用户反馈、赛后试点场地和行业伙伴'],
    ], [2400,3300,3660], font_size=9.5)
    source(doc, '官方来源：makerinchina.hk（2026 香港站主页，访问日期 2026-08-20）。历史公开赛事资料也显示，评审重点长期围绕技术/产品、商业模式、行业市场和团队；具体赛制以组委会最新通知为准。')
    heading(doc, '建议申报主赛道', 2)
    bullet(doc, '主赛道：智慧生活与智慧出行。Careloop 的核心是社区餐饮服务和长者可达性。')
    bullet(doc, '技术标签：人工智能与大数据。仅使用“语音意图解析、订单约束、置信度、人工兜底”等可以现场证明的表述。')
    bullet(doc, '辅助标签：智能装备与机器人。将自动烹饪、传送和回收定位为未来可插拔的履约模块，不宣称当前已完成。')
    heading(doc, '参赛叙事的重心变化', 2)
    table(doc, ['旧版叙事', '新版叙事'], [
        ['我们要建设大湾区自动化社区餐饮生态', '我们已经做出可运行的适老点餐与履约原型'],
        ['机器负责重活，生态负责关怀', '软件先验证需求，自动化按订单密度逐步接入'],
        ['10 个社区、180 天、RMB 500,000', '提交前做可重复的远程 Demo 验证；赛后再做香港 90 日试点'],
        ['复购率和回本期作为成果', '复购率和单位经济作为赛后待验证的试点 Gate'],
    ], [4680,4680], font_size=10)

    page_break(doc)
    heading(doc, '二、TeaVita AI 样本带来的可借鉴结构', 1)
    add_para(doc, 'TeaVita AI 的已提交材料有一个明显特点：它把“市场机会、技术依据、产业合作、实际 traction、护城河、路线图、财务和团队”分别变成了可扫描的证据页。Careloop 应借鉴这种证据结构，但不能复制它的真实收入、专利、供应链或合作协议。')
    table(doc, ['TeaVita 样本中的证据页', 'Careloop 的对应模块', 'Careloop 当前证据等级'], [
        ['Market Disconnect / Opportunity', '香港长者餐饮、数字鸿沟、社区照顾的公开资料与用户访谈', '需补：外部数据 + 访谈'],
        ['Scientific Validation', '语音意图解析、订单确认、状态机、错误转人工', '已有：工程证据'],
        ['Strategic Alliance', '香港社区/食堂/物业/社福合作方', '需补：意向或试点联系人'],
        ['Market Validation & Traction', 'Demo 使用路径、可用性测试、用户任务完成率', '需补：小规模测试'],
        ['Competitive Moat', '安全确认、人机协作、运营流程数据、场景关系', '部分成立：需避免过早使用 Data Moat'],
        ['Roadmap / Finance / Risk', '赛后香港 90 日试点、单位经济假设、食品/隐私/运营风险', '可写：必须标注假设'],
        ['Team', '软件、语音、UI、运营/长者服务能力与缺口', '需补：角色与外部顾问'],
    ], [2550,3900,2910], font_size=9.2)
    callout(doc, '关键原则', '没有数据时，不要用更多形容词填空；应把“我们还不知道什么”写成一个设计良好的验证计划。', fill='FFF4CC', label_color=GOLD)
    heading(doc, 'Careloop 可使用的“早期项目证据”', 2)
    for x in [
        '产品证据：新版 UI 的完整端到端演示，而不是只展示概念模型。',
        '工程证据：模块边界、API、状态机、异常路径、自动化测试结果。',
        '可用性证据：至少 5 位目标用户完成任务测试；记录完成时间、错误点和人工介入点。',
        '需求证据：访谈记录、原话、愿意试用/愿意付费/愿意合作的分层结果。',
        '落地证据：赛后试点场景清单、合作方画像、90 日排期和明确成功标准；提交前不宣称已落地。',
    ]:
        bullet(doc, x)

    page_break(doc)
    heading(doc, '三、旧版 Slides 的主要问题与修改原则', 1)
    table(doc, ['问题', '表现', '处理'], [
        ['定位过大', '社区餐饮生态、自动烹饪、冷链、社交、健康、自动回收同时出现', '收窄到“适老点餐与履约操作系统”；其他作为 roadmap'],
        ['地域错位', '香港站材料大量使用 Guangzhou/GBA 10 个社区', '改为 Hong Kong first，再讲大湾区复制'],
        ['证据混级', '计划、假设、演示和成果混在同一页', '每页标注：已实现 / 外部依据 / 待验证'],
        ['指标空转', '7 日复购、30 日复购、单机 15 单等没有基线', '改为试点指标和决策 Gate'],
        ['财务不一致', 'RMB 250,000、RMB 500,000、8 个月回本并存', '暂不使用精确财务结论；先建立单位经济模板'],
        ['Demo 位置偏后', '真实 UI 与系统闭环没有成为早期主证据', '第 4-5 页就展示新版 UI 和 90 秒 Demo'],
        ['团队页像演讲分工表', '主要展示谁负责哪一部分', '改为“能力—证据—缺口—补齐计划”'],
    ], [1800,3780,3780], font_size=9.3)
    heading(doc, '必须从旧版删除或降级的表达', 2)
    bullet(doc, '“10 个大湾区社区已经试点”——改为“赛后目标：在香港完成 1 个 90 日验证”。')
    bullet(doc, '“7 日/30 日复购”——改为“赛后试点期间将追踪 7 日/30 日复购；提交前不虚构结果”。')
    bullet(doc, '“约 8 个月回本”——改为“回本期取决于订单密度、人工、场地和食材成本，待试点验证”。')
    bullet(doc, '“数据护城河已经形成”——改为“订单偏好和运营数据将成为长期可积累资产”。')
    bullet(doc, '“AI 自动烹饪生态”——改为“当前是软件闭环；自动化硬件是后续履约模块”。')

    page_break(doc)
    heading(doc, '四、基于当前 project 的真实产品证据', 1)
    add_para(doc, '当前 project 的价值不是“已经有市场”，而是已经有一个可演示、可测试、可扩展的服务闭环。新版 Slides 应以此作为最核心的证据。')
    table(doc, ['模块', '当前可演示内容', '路演说法'], [
        ['顾客点餐', '菜单、菜品、口味、餐桌选择、订单预览', '长者可以用低认知负担完成点餐'],
        ['语音点餐', '语音/文字输入，菜品和餐桌意图解析，置信度与候选', '系统不是盲目执行，而是先理解再确认'],
        ['安全确认', '订单必须经过预览和用户确认后提交', '避免误识别直接变成错误订单'],
        ['为他人点餐', '家属/代理人模式和待确认订单', '把“远程关怀”落到可执行订单'],
        ['工作人员协助', '语音失败或未知菜品时转人工', '适老产品必须有人机协作兜底'],
        ['后厨履约', '后厨任务状态推进，订单与厨房任务关联', '让前端承诺和后端执行连接起来'],
        ['配送模拟', '传送带任务、合法桌号、配送状态', '展示未来智能履约接口，但不虚构真实硬件'],
        ['工程质量', '2 个测试文件、9 个测试通过；build/lint 已通过', '原型不是静态页面，而是可验证软件系统'],
    ], [1800,4050,3510], font_size=9.2)
    source(doc, '内部验证：2026-08-20 在当前 project 运行 npm test、npm run build、npm run lint；测试 2 files / 9 passed，build 与 lint 通过。该结果属于工程验证，不属于用户或商业 traction。')
    heading(doc, '技术闭环页建议文案', 2)
    add_para(doc, '用户说一句话 → 系统提取菜品、数量、口味和餐桌 → 展示订单预览与置信度 → 用户确认 → 后厨任务生成 → 状态实时更新 → 识别失败时转工作人员。', size=13, color=BLUE_DARK, bold=True)
    heading(doc, '必须主动披露的 Demo 边界', 2)
    bullet(doc, '支付目前为模拟适配器，不应在 Slides 中写成已接入真实支付。')
    bullet(doc, 'AICAN 和传送带为 Demo adapter，不应写成已经连接物理设备。')
    bullet(doc, '浏览器语音识别不可用时，当前系统提供文字备用和工作人员帮助路径。')
    bullet(doc, '当前 UI 主要验证交互和服务流程，食品安全、真实厨房和配送仍需赛后现场试点。')

    page_break(doc)
    heading(doc, '五、新版 12-14 页主路演建议', 1)
    table(doc, ['页码', '建议标题', '核心内容', '证据状态'], [
        ['1', 'Careloop：让社区餐饮更容易被长者使用', '一句话定位、当前原型截图、香港站赛道', '已实现 + 定位'],
        ['2', '问题不是没有饭，而是点餐、等待和照顾都不连续', '香港场景、长者/家属/食堂三方痛点', '外部资料 + 访谈待补'],
        ['3', '现有方案为什么仍然让老人独自承担', '传统食堂、纯 App、家属代点的缺口', '分析性判断'],
        ['4', '新版 Careloop UI：长者、家属、工作人员共用一条闭环', '新版 UI 截图与角色切换', '已实现'],
        ['5', '90 秒 Demo：从一句话到一份可履约订单', '现场演示脚本和订单状态变化', '已实现'],
        ['6', '技术闭环：理解、确认、履约、兜底', '语音意图、置信度、订单状态机、人工协助', '已实现/工程验证'],
        ['7', '我们今天已经验证了什么', '功能清单、测试结果、边界披露', '已实现'],
        ['8', '我们还没有验证什么', '付费、复购、真实场地、真实运营成本', '诚实披露'],
        ['9', 'Hong Kong First：赛后 90 日验证路线', '当前不做线下承诺；展示赛后场景、用户、合作方、排期和成功标准', '赛后计划'],
        ['10', '商业模式：从软件闭环开始，再接入自动化', '食堂/物业、家属、社区机构的价值和收费假设', '待验证假设'],
        ['11', '风险与合规：先做低风险可控的服务', '隐私、食品安全、人工兜底、硬件依赖', '风险计划'],
        ['12', '团队：我们能做什么，还缺什么', '技术能力、运营缺口、顾问/合作方需求', '团队事实'],
        ['13', '我们希望从香港站获得什么', '远程反馈、导师/投资者、赛后试点场地、食堂/物业伙伴、技术对接', '明确 Ask'],
        ['14', '从香港验证，到大湾区复制', '90 日后决策 Gate 与长期愿景', '路线图'],
    ], [720,2200,4500,1940], font_size=8.8)
    callout(doc, '删减原则', '旧版第 17-21 页的规模化试点、财务和补贴内容不必全部删除，可移入附录，并统一标注“项目假设 / 待验证”。', fill='FFF4CC', label_color=GOLD)

    page_break(doc)
    heading(doc, '六、90 秒现场 Demo 讲稿', 1)
    add_para(doc, 'Demo 目标不是展示所有功能，而是证明一个完整、可信、可失败恢复的订单闭环。建议使用一个长者角色、一个家属角色和一个工作人员兜底场景。')
    table(doc, ['时间', '动作', '屏幕必须出现', '评委应理解'], [
        ['0-10 秒', '介绍李阿姨：视力下降、独居、不会使用复杂 App', '长者模式首页', '问题是使用门槛，不是没有菜单'],
        ['10-25 秒', '说：“我要一份少盐红烧肉、一碗软一点的小米粥，送到 A12 桌。”', '语音输入/识别结果', '系统理解自然语言'],
        ['25-40 秒', '展示菜品、数量、口味、餐桌和置信度', '订单预览', '系统不会未经确认直接下单'],
        ['40-50 秒', '点击确认', '订单状态：已确认/制作中', '前端订单和后厨任务连起来'],
        ['50-65 秒', '切换后厨/配送视图', '厨房任务、配送状态、桌号', '存在可执行的运营流程'],
        ['65-80 秒', '故意输入不存在的菜品或模糊语句', '候选结果/工作人员协助', '失败时有人机协作兜底'],
        ['80-90 秒', '切换家属代点模式', '家属代点/待确认订单', 'Careloop 不只是一个老人 App'],
    ], [1100,2800,2800,2660], font_size=9.0)
    heading(doc, 'Demo 现场不要做的事', 2)
    bullet(doc, '不要花大量时间讲自动烹饪机器的内部结构，当前项目没有真实设备证据。')
    bullet(doc, '不要把模拟传送带说成真实硬件演示。')
    bullet(doc, '不要只展示漂亮 UI 而跳过错误处理和用户确认。')
    bullet(doc, '不要在现场依赖不可控的网络或浏览器语音；准备预录视频和文字输入备份。')

    page_break(doc)
    heading(doc, '七、内容补强方案：没有市场数据时如何让 Slides 不空', 1)
    add_para(doc, '建议把证据分成四个层次，并在每一页右下角用小标签标记：VALIDATED、EXTERNAL SOURCE、DEMO ASSUMPTION、TO VALIDATE。')
    table(doc, ['证据层次', '可以使用的内容', '不可越界的表述'], [
        ['A. 已验证工程事实', '功能、接口、测试、状态机、错误路径', '不能称为用户验证或市场验证'],
        ['B. 外部公开资料', '香港人口、长者服务、数字鸿沟、餐饮运营资料', '不能把大市场等同于 Careloop 需求已成立'],
        ['C. 轻量用户研究', '访谈、可用性测试、任务完成率、意向登记', '不能把“喜欢概念”写成“愿意付费”'],
    ['D. 赛后试点计划', '90 日排期、成功标准、决策 Gate、合作方需求', '不能写成已经完成的 pilot/traction'],
    ], [1900,4100,3360], font_size=9.3)
    heading(doc, '当前时间节点可完成的最小证据包', 2)
    numbered(doc, '整理当前 project 的功能清单、API 路径、状态机和测试结果，形成 1 页工程验证表。')
    numbered(doc, '录制一条 90 秒完整 Demo 视频，并准备文字输入和预录视频作为语音/网络故障备份。')
    numbered(doc, '邀请团队外的同学、导师或远程体验者完成同一套 UI 任务；若无法招募，必须将其标记为“未完成”，不能写成用户验证。')
    numbered(doc, '把所有未来数字改成目标或假设，并在页脚加入“赛后香港试点验证”的注释。')
    numbered(doc, '将长者、家属、食堂和物业访谈列为赛后第一阶段研究，不把未发生的访谈写入本次提交。')
    heading(doc, '推荐在 Slides 中呈现的轻量数据', 2)
    table(doc, ['指标', '记录方式', '可写成什么'], [
        ['任务完成率', '完成/未完成人数', '“在内部可用性测试中，X/Y 人完成订单预览”'],
        ['平均完成时间', '从开始到确认的秒数', '“目标是将首次点餐控制在 X 分钟内”'],
        ['人工接管率', '发生错误或求助的次数', '“人工兜底仍是设计的一部分”'],
        ['最常见错误', '分类记录', '“主要风险来自菜品别名/口音/桌号确认”'],
        ['意愿登记', '愿意试用/愿意付费/愿意合作分层', '“兴趣不等于付费，下一步在试点验证”'],
    ], [2300,3000,4060], font_size=9.2)

    page_break(doc)
    heading(doc, '八、赛后香港 90 日试点路线（不是当前承诺）', 1)
    add_para(doc, '由于当前提交节点不适合进行线下运营，本节只作为赛后落地路线图。当前路演不应声称已经完成试点，也不应把合作方、用户和复购数据写成既成事实。')
    table(doc, ['阶段', '周期', '目标', '输出'], [
        ['准备', '赛后第 1-2 周', '确定 1 个社区场景、菜单、桌号和工作人员角色', '试点协议、用户招募、隐私告知'],
        ['基线', '第 3-4 周', '记录传统点餐时间、错误、人工成本和订单峰值', 'Baseline dashboard'],
        ['小规模运行', '第 5-8 周', '让 30-50 位长者/家属使用点餐和代点', '任务完成率、接管率、满意度'],
        ['复购验证', '第 9-12 周', '观察再次使用、家属代点和食堂运营负担', '7 日/30 日复购、单位经济初版'],
    ], [1500,1700,3200,2960], font_size=9.2)
    heading(doc, '试点成功标准（建议写成 Gate，而不是承诺结果）', 2)
    table(doc, ['优先级', '指标', '建议判断方式'], [
        ['P1 需求', '长者愿意完成订单；家属愿意代点', '达到预设任务完成率和首次使用转化'],
        ['P2 运营', '订单能被工作人员接住并完成履约', '高峰期不出现不可解释的订单丢失'],
        ['P3 经济', '每餐毛利、人工和场地成本可测', '得到单站点盈亏平衡点，而非直接承诺回本'],
        ['P4 留存', '7 日/30 日再次使用', '确认是一次性新鲜感还是日常习惯'],
    ], [1500,3500,4360], font_size=9.2)
    heading(doc, '赛后需要的合作方', 2)
    bullet(doc, '一个可以接触长者的社区/社福/长者中心。')
    bullet(doc, '一个愿意提供菜单、桌号和工作人员流程的社区食堂或餐饮方。')
    bullet(doc, '一个可以讨论场地、住户触达和长期运营的物业或社区运营方。')
    bullet(doc, '一位长者服务或食品安全顾问，帮助避免把 Demo 直接当成可上线餐饮系统。')

    page_break(doc)
    heading(doc, '九、商业模式与财务表达', 1)
    heading(doc, '当前阶段的商业模式', 2)
    table(doc, ['付费方', '价值', '可测试的收费方式'], [
        ['社区食堂/物业', '减少人工沟通、统一订单和履约状态、提升住户服务', '站点月费 + 订单服务费'],
        ['家属', '远程代点、查看订单状态、形成照顾记录', '免费基础功能 + 家属增值服务'],
        ['社区机构/社福项目', '提升长者服务可达性和运营可视化', '项目采购/服务合同'],
        ['后续硬件伙伴', '把高频、稳定的流程自动化', '设备接入/运营服务费'],
    ], [2100,3900,3360], font_size=9.2)
    heading(doc, '财务页先写“模型”，不要写“结论”', 2)
    add_para(doc, '建议在内部版本建立以下变量，所有数字都从香港真实报价、场地和试点数据进入：')
    table(doc, ['变量', '需要收集的事实', '当前状态'], [
        ['每餐售价', '社区食堂菜单、长者可接受价格', '待访谈/报价'],
        ['食材成本率', '菜单、供应商和损耗', '待获取'],
        ['人工成本', '点餐协助、厨房、配送、清洁', '待测量'],
        ['场地成本', '物业/食堂合作模式', '待谈判'],
        ['软件收费', '站点月费或按单收费的接受度', '待验证'],
        ['盈亏平衡订单量', '由以上变量计算', '不要提前写死'],
    ], [2100,4360,2900], font_size=9.2)
    callout(doc, '建议删掉', 'RMB 500,000、10 个社区和约 8 个月回本，除非能提供一套清晰的香港成本表、供应商报价和订单假设。否则应放入“扩张情景”，不是当前事实。', fill='FFF4CC', label_color=GOLD)

    page_break(doc)
    heading(doc, '十、风险、隐私和合规表达', 1)
    table(doc, ['风险', '当前风险点', 'Demo/试点控制'], [
        ['语音误识别', '菜品、数量、口味、桌号可能错误', '预览确认、置信度、候选、人工接管'],
        ['老人不会用', '数字界面增加挫败感', '大按钮、语音/文字双通道、社区工作人员协助'],
        ['订单丢失/状态错误', '前台和后厨不同步', '订单状态机、事件记录、人工查询'],
        ['隐私', '家属代点和长者偏好涉及个人数据', '最小化收集、权限分层、隐私告知、删除/过期机制'],
        ['食品安全', '软件不能替代厨房监管', '试点使用现有合规食堂和明确责任边界'],
        ['硬件依赖', '自动烹饪或传送设备尚未真实接入', '人工流程兜底；硬件作为后续接口'],
        ['商业不成立', '用户喜欢但不愿意付费', '设置停止/优化 Gate，不把情感反馈当收入'],
    ], [1800,3900,3660], font_size=9.2)
    heading(doc, '必须使用的诚实表述', 2)
    add_para(doc, '“当前版本用于验证点餐与履约服务流程，不构成医疗建议，也不替代食品安全和社区照顾专业人员。任何真实部署均需由场地方、食品服务方和相关顾问共同确认。”', size=11.5, color=RED, bold=True, after=8)

    page_break(doc)
    heading(doc, '十一、团队页与合作诉求', 1)
    heading(doc, '团队页改法', 2)
    table(doc, ['能力', '必须展示的证据', '若当前不足'], [
        ['产品/工程', '当前 project、API、状态机、测试和 UI', '已有，可直接展示'],
        ['语音/AI', '意图解析、置信度、候选和人工兜底', '已有原型；避免过度称为大模型能力'],
        ['长者服务', '访谈、顾问、社区资源', '需要补齐外部顾问/合作方'],
        ['餐饮运营', '食堂流程、菜单、厨房、食品安全', '需要试点方共同设计'],
        ['商业化', '香港首站模型、收费假设、扩张路线', '由试点结果驱动，不提前承诺规模'],
    ], [1800,4300,3260], font_size=9.2)
    heading(doc, '结尾的 Ask 建议', 2)
    add_para(doc, '我们当前不把线下试点包装成已发生的成果，而是希望通过香港站获得远程反馈、导师/投资者指导和赛后试点连接：一个社区/长者中心、一个食堂或物业合作方、30-50 位目标用户，以及一位能够指导长者服务和食品安全的伙伴。赛后 90 日，我们再用订单完成、人工介入、复购和单位经济决定是否扩张。', size=12.5, color=BLUE_DARK, bold=True)
    table(doc, ['Ask 类型', '具体请求'], [
        ['赛后试点场景', '协助对接一个社区、长者中心、屋苑会所或社区食堂'],
        ['当前远程反馈', '邀请评审、导师、投资者或行业人士体验新版 UI 并指出阻碍'],
        ['用户触达', '协助招募长者、家属和工作人员'],
        ['行业合作', '提供菜单、运营、食品安全和场地流程建议'],
        ['技术合作', '语音、支付、设备接入或数据安全支持'],
        ['资本/政策', '支持赛后 90 日验证，而不是直接支持尚未验证的十站扩张'],
    ], [2000,7360], font_size=9.5)

    page_break(doc)
    heading(doc, '十二、实施清单与排练标准', 1)
    heading(doc, '提交前必须补齐', 2)
    for x in [
        '确认香港非上市公司/团队主体、报名分类、联系人和公司资料。',
        '把新版 UI 截图和真实 Demo 视频纳入 Slides；准备无网络备份。',
        '若无法在提交前完成真实访谈，不要补造；改为补充可复现 Demo、工程测试和赛后用户研究计划。',
        '把所有数字标成“已验证 / 外部资料 / 假设 / 待验证”。',
        '删掉或降级没有证据的收入、复购、回本、专利、试点和自动化表述。',
        '统一 Careloop、暖桌、Careloop Dining 的品牌名称。',
        '修正英文拼写、双语术语和页面中重复的 AICAN Visual Reference 标识。',
        '准备 90 秒 Demo、7 分钟主讲、5 分钟答辩和 2 分钟故障备份版本；不把线下试点作为当前承诺。',
    ]:
        bullet(doc, x)
    heading(doc, '答辩高概率问题', 2)
    table(doc, ['问题', '建议回答方向'], [
        ['你们真正自主研发的是什么？', '当前自主完成的是适老点餐与履约软件闭环；硬件不是当前核心成果，后续通过接口和合作方接入。'],
        ['有没有真实老人使用？', '目前没有规模化线下试点；我们完成了工程原型，当前用可复现 Demo 和工程验证展示能力，赛后再通过香港 90 日试点验证可用性、付费和复购。'],
        ['为什么不是普通点餐 App？', '核心差异是老人场景的语音/一键、确认、人工兜底、家属代点和后厨履约被设计成一条闭环。'],
        ['AI 在哪里？', '语音转订单意图、菜品约束、置信度、候选和异常转人工；不把未接入的模型和硬件包装成现成果。'],
        ['如何赚钱？', '先验证食堂/物业的站点服务价值，再测试月费、订单服务费和机构项目采购。'],
        ['为什么从香港开始？', '香港适合验证高密度社区、长者服务和家属代点；验证后再复制到大湾区。'],
    ], [3000,6360], font_size=9.3)
    heading(doc, '最终判断标准', 2)
    callout(doc, '路演成功', '评委在演示后能准确复述：Careloop 为谁解决什么问题、今天已经做到了什么、还缺什么证据、希望在香港获得什么合作。', fill=PALE_TEAL, label_color=TEAL)

    page_break(doc)
    heading(doc, '附录 A：建议使用的证据标签', 1)
    table(doc, ['标签', '含义', '例子'], [
        ['VALIDATED / 已验证', '当前 project 或内部工程测试已完成', '2 个测试文件、9 个测试通过'],
        ['EXTERNAL SOURCE / 外部资料', '来自官网、政府、研究或公开报告', '赛事赛道、时间表、公开人口/行业资料'],
        ['DEMO ASSUMPTION / Demo 假设', '为了展示流程的模拟内容', '传送带、自动烹饪、桌号地图'],
        ['TO VALIDATE / 待验证', '需要远程用户反馈或赛后香港场地验证', '付费、复购、订单密度、人工成本'],
    ], [2100,3600,3660], font_size=9.5)
    heading(doc, '附录 B：参考资料', 1)
    refs = [
        '1. Maker in China Hong Kong Chapter 2026 官网：https://makerinchina.hk/（赛事时间、报名对象、科技领域、奖励与对接信息；访问日期 2026-08-20）',
        '2. Hong Kong Digital Policy Office：Maker in China 历届赛事总结：https://www.digitalpolicy.gov.hk/sc/our_work/digital_infrastructure/mainland/maker_in_china/（历届项目、报名规模、产业对接和晋级信息）',
        '3. Maker in China Hong Kong 2025 Contest Review：https://makerinchina.hk/2025/（2025 十强与赛后信息）',
        '4. Careloop 当前 project README：/Users/cj/careloop/README.md（当前 Demo 范围与适配器边界）',
        '5. Careloop 当前 project：/Users/cj/careloop/src/ui/index.html、/Users/cj/careloop/src/api/orders.ts、/Users/cj/careloop/packages/domain/src/voice/index.ts（UI、API、语音意图和状态流转）',
        '6. 旧版 Slides：/Users/cj/careloop/slides/Group 5-暖桌Careloop.pdf（旧版内容基线）',
        '7. TeaVita AI 已提交样本：/Users/cj/careloop/slides/TeaVita AI--Maker in China.pdf（同赛事表达结构参考，不作为 Careloop 事实来源）',
    ]
    for ref in refs:
        bullet(doc, ref)
    add_para(doc, '文档版本：v1.0 | 编制日期：2026-08-20 | 状态：内部改稿基线', size=9, color=GRAY, italic=True, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build()
